import os
import logging
import requests
from bs4 import BeautifulSoup
from src.repository.vector_db import get_vector_store

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Minimal Document dataclass (replaces langchain_core.documents.Document)
# ---------------------------------------------------------------------------

class Document:
    """Lightweight stand-in for langchain Document — page_content + metadata."""
    __slots__ = ("page_content", "metadata")

    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}


# ---------------------------------------------------------------------------
# Simple recursive-character text splitter (replaces langchain_text_splitters)
# ---------------------------------------------------------------------------

def _split_text(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> list:
    """
    Split *text* into overlapping chunks no larger than *chunk_size*.
    Strategy: paragraph → line → word (mirrors RecursiveCharacterTextSplitter).
    """
    if not text or not text.strip():
        return []

    def _merge(parts: list, sep: str) -> list:
        """Merge parts into chunks, respecting chunk_size and chunk_overlap."""
        chunks = []
        current = ""
        for part in parts:
            part = part.strip()
            if not part:
                continue
            candidate = (current + sep + part).strip() if current else part
            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                    # Carry overlap into the next chunk
                    overlap = current[-chunk_overlap:].strip() if chunk_overlap else ""
                    current = (overlap + sep + part).strip() if overlap else part
                else:
                    # single part already too large — split at word level
                    words = part.split()
                    for word in words:
                        c2 = (current + " " + word).strip() if current else word
                        if len(c2) <= chunk_size:
                            current = c2
                        else:
                            if current:
                                chunks.append(current)
                            overlap = current[-chunk_overlap:].strip() if chunk_overlap and current else ""
                            current = (overlap + " " + word).strip() if overlap else word
        if current:
            chunks.append(current)
        return chunks

    # Try paragraph split first
    paragraphs = text.split("\n\n")
    if len(paragraphs) > 1:
        chunks = _merge(paragraphs, "\n\n")
    else:
        # Fall back to line split
        lines = text.split("\n")
        if len(lines) > 1:
            chunks = _merge(lines, "\n")
        else:
            chunks = _merge([text], " ")

    # Second pass: any chunk still over chunk_size gets word-split
    final = []
    for chunk in chunks:
        if len(chunk) <= chunk_size:
            final.append(chunk)
        else:
            words = chunk.split()
            current = ""
            for word in words:
                c = (current + " " + word).strip() if current else word
                if len(c) <= chunk_size:
                    current = c
                else:
                    if current:
                        final.append(current)
                    overlap = current[-chunk_overlap:].strip() if chunk_overlap and current else ""
                    current = (overlap + " " + word).strip() if overlap else word
            if current:
                final.append(current)

    return [c for c in final if c.strip()]


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------

class DocumentIngestionService:
    """Ingest documents into the vector store for RAG.
    Uses Google Gemini embeddings API — no PyTorch, no langchain, no local model.
    """

    CHUNK_SIZE = 500
    CHUNK_OVERLAP = 100

    def __init__(self):
        self.vector_store = get_vector_store()

    def _split(self, text: str) -> list:
        return _split_text(text, self.CHUNK_SIZE, self.CHUNK_OVERLAP)

    def _add_chunks(self, chunks: list, title: str, extra_meta: dict = None) -> bool:
        """Wrap text chunks into Documents and add to the vector store."""
        if not chunks or not self.vector_store:
            return False
        meta = {"source": title}
        if extra_meta:
            meta.update(extra_meta)
        docs = [Document(page_content=c, metadata=meta) for c in chunks]
        self.vector_store.add_documents(docs)
        logger.info(f"Ingested {len(docs)} chunks from '{title}'")
        return True

    # ------------------------------------------------------------------
    # File ingestors
    # ------------------------------------------------------------------

    def ingest_text_file(self, file_path: str, title: str) -> bool:
        """Ingest a plain text file."""
        try:
            if not os.path.exists(file_path):
                logger.warning(f"File not found: {file_path}")
                return False
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            if not content.strip():
                logger.warning(f"File is empty: {file_path}")
                return False
            return self._add_chunks(self._split(content), title, {"file_path": file_path})
        except Exception as e:
            logger.error(f"Error ingesting text file {file_path}: {e}")
            return False

    def ingest_pdf_file(self, file_path: str, title: str) -> bool:
        """Ingest a PDF file using pypdf."""
        try:
            if not os.path.exists(file_path):
                logger.warning(f"PDF not found: {file_path}")
                return False
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            all_chunks = []
            for page in reader.pages:
                text = page.extract_text() or ""
                if text.strip():
                    all_chunks.extend(self._split(text))
            if not all_chunks:
                logger.warning(f"No content extracted from PDF: {file_path}")
                return False
            return self._add_chunks(all_chunks, title, {"file_path": file_path})
        except Exception as e:
            logger.error(f"Error ingesting PDF {file_path}: {e}")
            return False

    def ingest_docx_file(self, file_path: str, title: str) -> bool:
        """Ingest a DOCX file."""
        try:
            if not os.path.exists(file_path):
                logger.warning(f"DOCX not found: {file_path}")
                return False
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            content = "\n".join(p.text for p in doc.paragraphs)
            if not content.strip():
                logger.warning(f"No content extracted from DOCX: {file_path}")
                return False
            return self._add_chunks(self._split(content), title, {"file_path": file_path})
        except Exception as e:
            logger.error(f"Error ingesting DOCX {file_path}: {e}")
            return False

    # ------------------------------------------------------------------
    # Web scraper
    # ------------------------------------------------------------------

    def scrape_webpage(self, url: str) -> str:
        """Scrape and extract plain text from a webpage."""
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            response = requests.get(url, timeout=10, headers=headers)
            response.raise_for_status()
            soup = BeautifulSoup(response.content, "html.parser")
            for tag in soup(["script", "style", "meta", "noscript"]):
                tag.decompose()
            lines = (line.strip() for line in soup.get_text().splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return "\n".join(c for c in chunks if c)
        except Exception as e:
            logger.error(f"Error scraping {url}: {e}")
            return ""

    def ingest_link(self, url: str, title: str) -> bool:
        """Fetch and ingest content from a URL."""
        try:
            content = self.scrape_webpage(url)
            if not content.strip():
                logger.warning(f"No content from URL: {url}")
                return False
            chunks = self._split(content)
            if not chunks:
                return False
            return self._add_chunks(chunks, title, {"url": url, "type": "link"})
        except Exception as e:
            logger.error(f"Error ingesting link {url}: {e}")
            return False

    # ------------------------------------------------------------------
    # Router
    # ------------------------------------------------------------------

    def ingest_document(self, file_path: str, file_name: str, title: str, doc_type: str) -> bool:
        """Route ingestion based on document type."""
        if doc_type == "link":
            return self.ingest_link(file_path, title)
        if file_name.endswith(".pdf"):
            return self.ingest_pdf_file(file_path, title)
        elif file_name.endswith(".txt"):
            return self.ingest_text_file(file_path, title)
        elif file_name.endswith(".docx"):
            return self.ingest_docx_file(file_path, title)
        else:
            logger.warning(f"Unsupported file type: {file_name}")
            return False
