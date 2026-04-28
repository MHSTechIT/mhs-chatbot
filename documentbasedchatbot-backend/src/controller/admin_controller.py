import os
import io
import logging
import uuid as _uuid_module
from datetime import datetime
from fastapi import APIRouter, Depends, File, UploadFile, Form, HTTPException
from pydantic import BaseModel
from src.repository.admin_repo import get_admin_repository
from src.repository.enrollment_repo import EnrollmentRepository
from src.database import SessionLocal
from src.middleware.admin_auth import verify_admin_key
# DocumentIngestionService deferred to avoid importing torch/HuggingFace at startup

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/admin", tags=["Admin"], dependencies=[Depends(verify_admin_key)])

# All services lazily initialized — no DB/ML connections at import time
ingestion_service = None

def _get_admin_repo():
    return get_admin_repository()

def _get_ingestion_service():
    global ingestion_service
    if ingestion_service is None:
        from src.services.DocumentIngestionService import DocumentIngestionService
        try:
            ingestion_service = DocumentIngestionService()
        except Exception as e:
            logger.warning(f"Document ingestion service not available: {str(e)}")
    return ingestion_service


class LinkRequest(BaseModel):
    title: str
    url: str


class DocumentResponse(BaseModel):
    id: str
    title: str
    type: str
    url: str = None
    file_name: str = None
    uploaded_at: str


class EnrollmentRequest(BaseModel):
    name: str
    phone: str
    sugar_level: str = None


@router.post("/upload")
async def upload_document(file: UploadFile = File(...), title: str = Form(...)):
    """Upload a document (PDF, TXT, DOCX) - saves content to DATABASE"""
    if not title.strip():
        raise HTTPException(status_code=400, detail="Title required")

    db = None
    try:
        # Save file
        upload_dir = "uploads"
        os.makedirs(upload_dir, exist_ok=True)

        file_path = os.path.join(upload_dir, file.filename)
        file_content = await file.read()

        # Extract text content based on file type
        text_content = ""
        fname_lower = file.filename.lower()
        if fname_lower.endswith('.pdf'):
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_content))
                pages_text = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text_content = "\n".join(pages_text)
                logger.info(f"✅ Extracted PDF text: {len(text_content)} characters from {len(reader.pages)} pages")
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        elif fname_lower.endswith('.docx'):
            try:
                from docx import Document as DocxDoc
                doc_obj = DocxDoc(io.BytesIO(file_content))
                text_content = "\n".join(p.text for p in doc_obj.paragraphs if p.text.strip())
                logger.info(f"✅ Extracted DOCX text: {len(text_content)} characters")
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
        else:
            # Plain text — try UTF-8 then latin-1
            try:
                text_content = file_content.decode('utf-8')
                logger.info(f"✅ Read text content: {len(text_content)} characters")
            except Exception:
                try:
                    text_content = file_content.decode('latin-1')
                    logger.info(f"✅ Read text (latin-1): {len(text_content)} characters")
                except Exception:
                    logger.warning("Could not decode file content")

        # Save file to disk
        with open(file_path, "wb") as f:
            f.write(file_content)

        logger.info(f"✅ File uploaded: {file.filename}")

        # Save to database with content
        db = SessionLocal()
        from src.models.document import Document

        doc_uuid = _uuid_module.uuid4()   # UUID object for Supabase uuid column
        doc_id = str(doc_uuid)            # String for in-memory dict / response
        new_doc = Document(
            id=doc_uuid,
            title=title.strip(),
            type="document",
            file_name=file.filename,
            file_path=file_path,
            content=text_content if text_content else None  # Store content in database
        )
        db.add(new_doc)
        db.commit()

        # Also update admin_repo for consistency
        _get_admin_repo().documents[doc_id] = {
            "id": doc_id,
            "title": title.strip(),
            "type": "document",
            "file_name": file.filename,
            "file_path": file_path,
            "content": text_content,
            "url": None,
            "uploaded_at": datetime.now().isoformat()
        }
        _get_admin_repo().save_documents()

        logger.info(f"✅ Document saved to database: {doc_id}")

        # Index document into vector store for RAG (optional)
        svc = _get_ingestion_service()
        if svc:
            try:
                success = svc.ingest_document(
                    file_path=file_path,
                    file_name=file.filename,
                    title=title,
                    doc_type="document"
                )
                if success:
                    logger.info(f"Document indexed into vector store: {title}")
                else:
                    logger.warning(f"Failed to index document into vector store: {title}")
            except Exception as e:
                logger.error(f"Error indexing document: {str(e)}")

        # ✅ Auto-delete uploaded file — content is already saved in DB and vector store
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"🧹 Deleted uploaded file after ingestion: {file_path}")
        except Exception as e:
            logger.warning(f"Could not delete uploaded file: {e}")

        return {"success": True, "id": doc_id, "message": "Document uploaded to database"}

    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")
    finally:
        if db:
            db.close()


@router.post("/add-link")
async def add_link(request: LinkRequest):
    """Add a reference link"""
    if not request.title.strip() or not request.url.strip():
        raise HTTPException(status_code=400, detail="Title and URL required")

    try:
        doc_id = _get_admin_repo().add_document(
            title=request.title,
            url=request.url,
            type="link"
        )

        # Index link into vector store for RAG
        svc = _get_ingestion_service()
        if svc:
            try:
                success = svc.ingest_link(
                    url=request.url,
                    title=request.title
                )
                if success:
                    logger.info(f"Link indexed into vector store: {request.title}")
                else:
                    logger.warning(f"Failed to index link into vector store: {request.title}")
            except Exception as e:
                logger.error(f"Error indexing link: {str(e)}")

        logger.info(f"Link added: {request.title}")
        return {"success": True, "id": doc_id, "message": "Link added"}

    except Exception as e:
        logger.error(f"Link error: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")


@router.get("/documents/content-status")
async def get_documents_content_status():
    """Diagnostic: show each document and how many characters of content are stored in DB"""
    try:
        _get_admin_repo().load_documents()
        result = []
        for doc in _get_admin_repo().documents.values():
            content_len = len(doc.get("content") or "")
            result.append({
                "id": doc["id"],
                "title": doc["title"],
                "type": doc["type"],
                "content_chars": content_len,
                "has_content": content_len > 100,
            })
        return {"success": True, "documents": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")


@router.get("/documents")
async def get_documents():
    """Get all uploaded documents and links"""
    try:
        # Force reload from file to get latest data
        _get_admin_repo().load_documents()
        documents = _get_admin_repo().get_all_documents()
        return {
            "success": True,
            "documents": documents,
            "count": len(documents)
        }
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document or link from database AND vector store"""
    try:
        doc = _get_admin_repo().get_document_by_id(doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Remove from vector store by matching the document title (source metadata)
        try:
            from src.repository.vector_db import get_vector_store
            vector_store = get_vector_store()
            if vector_store:
                count = vector_store.delete_by_source(doc['title'])
                logger.info(f"🧹 Removed {count} embeddings for '{doc['title']}' from vector store")
        except Exception as e:
            logger.error(f"Error removing from vector store: {e}")

        # Delete from database
        success = _get_admin_repo().delete_document(doc_id)
        if success:
            logger.info(f"Document deleted from database: {doc_id}")
            return {"success": True, "message": "Document deleted from database and vector store"}
        else:
            raise HTTPException(status_code=404, detail="Document not found in database")

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete error: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")


@router.get("/documents/search")
async def search_documents(query: str):
    """Search documents by title"""
    try:
        results = _get_admin_repo().search_documents(query)
        return {"success": True, "results": results}
    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")


@router.post("/submit-enrollment")
async def submit_enrollment(request: EnrollmentRequest):
    """Submit enrollment form with name, phone, and optional sugar level"""
    db = SessionLocal()
    try:
        if not request.name.strip() or not request.phone.strip():
            raise HTTPException(status_code=400, detail="Name and phone are required")

        enrollment_id = EnrollmentRepository.create_enrollment(
            db,
            name=request.name.strip(),
            phone=request.phone.strip(),
            sugar_level=request.sugar_level.strip() if request.sugar_level else None
        )

        logger.info(f"Enrollment submitted: {request.name} ({request.phone})")
        return {
            "success": True,
            "id": str(enrollment_id),
            "message": "Enrollment submitted successfully"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Enrollment submission error: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")
    finally:
        db.close()


# Pre-defined static messages to pre-record once and reuse forever
STATIC_AUDIO_TEXTS = {
    "welcome_ta": "சக்கரை நோய் பற்றி உங்களுக்கு ஏதாவது கேள்விகள் இருந்தால், தயங்காம கேளுங்கள்.",
    "enrollment_prompt_en": "For more details and personalized guidance, please fill in the form below — our team will contact you soon!",
    "enrollment_prompt_ta": "நம்ம course பத்தி more details தெரிஞ்சுக்கணும்னா, கீழே உள்ள form-ஐ fill பண்ணுங்க! எங்க team உங்களை விரைவில் contact பண்ணுவாங்க!",
    "post_enrollment_en": "Thank you! Our team will contact you soon and answer all your questions!",
    "post_enrollment_ta": "நன்றி! எங்க team விரைவில் உங்களை contact பண்ணி உங்க கேள்விகளுக்கு பதில் சொல்வாங்க!",
}


@router.post("/generate-static-audio")
async def generate_static_audio():
    """Generate and store pre-recorded audio for static messages (call ONCE — saves ElevenLabs credits)"""
    import base64
    import requests as req_lib
    from src.models.document import Document

    elevenlabs_api_key = os.getenv("ELEVENLABS_API_KEY")
    elevenlabs_voice_id = os.getenv("ELEVENLABS_VOICE_ID")
    if not elevenlabs_api_key or not elevenlabs_voice_id:
        raise HTTPException(status_code=500, detail="ElevenLabs credentials not configured")

    db = SessionLocal()
    results = {}
    try:
        for key, text in STATIC_AUDIO_TEXTS.items():
            # Skip if already generated
            existing = db.query(Document).filter(Document.title == f"static_audio:{key}").first()
            if existing and existing.content:
                results[key] = "already_exists"
                continue

            tts_url = f"https://api.elevenlabs.io/v1/text-to-speech/{elevenlabs_voice_id}"
            payload = {
                "text": text,
                "model_id": "eleven_turbo_v2_5",
                "voice_settings": {"stability": 0.5, "similarity_boost": 0.75, "style": 0.7, "use_speaker_boost": True},
            }
            headers = {"xi-api-key": elevenlabs_api_key, "Content-Type": "application/json"}
            resp = req_lib.post(tts_url, json=payload, headers=headers, timeout=30)
            if resp.status_code != 200:
                results[key] = f"error_{resp.status_code}"
                logger.error(f"ElevenLabs error for {key}: {resp.text[:200]}")
                continue

            audio_b64 = base64.b64encode(resp.content).decode("utf-8")

            if existing:
                existing.content = audio_b64
            else:
                doc_uuid = _uuid_module.uuid4()
                new_doc = Document(
                    id=doc_uuid,
                    title=f"static_audio:{key}",
                    type="static_audio",
                    file_name=f"{key}.mp3",
                    content=audio_b64,
                )
                db.add(new_doc)
            results[key] = "generated"
            logger.info(f"✅ Generated static audio: {key} ({len(audio_b64)} chars base64)")

        db.commit()
        return {"success": True, "results": results}
    except Exception as e:
        logger.error(f"Static audio generation error: {e}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")
    finally:
        db.close()


STATIC_AUDIO_CACHE_FILE = "data/static_audio.json"


def _load_static_audio_from_file() -> dict:
    """Load static audio from local JSON file (fallback when DB is unavailable)."""
    import json
    try:
        if os.path.exists(STATIC_AUDIO_CACHE_FILE):
            with open(STATIC_AUDIO_CACHE_FILE, "r") as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_static_audio_to_file(audio_map: dict):
    """Persist static audio to local JSON file so it survives DB outages."""
    import json
    try:
        os.makedirs("data", exist_ok=True)
        with open(STATIC_AUDIO_CACHE_FILE, "w") as f:
            json.dump(audio_map, f)
    except Exception as e:
        logger.warning(f"Could not save static audio to file: {e}")


@router.get("/static-audio")
async def get_static_audio():
    """Return all pre-recorded static audio as base64 — frontend caches these in localStorage.
    Reads from DB first; falls back to local JSON file if DB is unavailable."""
    from src.models.document import Document

    # Try DB first
    try:
        db = SessionLocal()
        try:
            docs = db.query(Document).filter(Document.type == "static_audio").all()
            if docs:
                audio_map = {doc.title.replace("static_audio:", ""): doc.content for doc in docs}
                _save_static_audio_to_file(audio_map)  # keep file in sync
                return {"success": True, "audio": audio_map}
        finally:
            db.close()
    except Exception as e:
        logger.warning(f"DB unavailable for static audio, falling back to file: {e}")

    # Fall back to file cache
    audio_map = _load_static_audio_from_file()
    return {"success": True, "audio": audio_map}


@router.get("/leads")
async def get_enrollment_leads(limit: int = 20, offset: int = 0):
    """Get enrollment leads with pagination"""
    db = SessionLocal()
    try:
        enrollments, total = EnrollmentRepository.get_enrollments_paginated(db, limit=limit, offset=offset)
        leads = [
            {
                "id": str(e.id),
                "name": e.name,
                "phone": e.phone,
                "sugar_level": e.sugar_level or "",
                "created_at": e.created_at.isoformat() if hasattr(e.created_at, 'isoformat') else str(e.created_at)
            }
            for e in enrollments
        ]
        return {"success": True, "leads": leads, "count": len(leads), "total": total}
    except Exception as e:
        logger.error(f"Error fetching leads: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to fetch leads.")
    finally:
        db.close()


@router.post("/clear-all")
async def clear_all_data():
    """DANGER: Clear ALL documents, links, and vector store data. Start fresh."""
    try:
        import shutil

        # Clear documents from memory cache
        _get_admin_repo().documents = {}

        # Clear documents from JSON file
        _get_admin_repo().save_documents()
        logger.info("Cleared all documents from JSON storage")

        # Verify by reloading from file
        _get_admin_repo().load_documents()
        if not _get_admin_repo().documents:
            logger.info("Verified: Documents cache is empty")

        # Delete uploaded files
        if os.path.exists("uploads"):
            shutil.rmtree("uploads")
            os.makedirs("uploads", exist_ok=True)
            logger.info("✅ Deleted uploads folder")

        # Clear all embeddings from vector store
        try:
            from src.repository.vector_db import get_vector_store
            vector_store = get_vector_store()
            if vector_store:
                vector_store.delete_collection()
                logger.info("✅ Cleared all embeddings from vector store")
        except Exception as e:
            logger.error(f"Error clearing vector store: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to clear vector store: {e}")

        logger.info("🔄 System cleared - ready to add new documents from scratch")
        return {
            "success": True,
            "message": "All data cleared successfully. Ready to start fresh."
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error clearing data: {str(e)}")
        raise HTTPException(status_code=500, detail="An internal error occurred. Please try again later.")
